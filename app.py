from flask import Flask, request, jsonify
from flask_cors import CORS, cross_origin
import PyPDF2
import io
import requests
import os
from dotenv import load_dotenv
import base64
from docx import Document
from PIL import Image
import pytesseract
import json # Import json for structured output

# --- Load environment variables ---
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# --- Retrieve API Key ---
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    print("\n" * 3)
    print("-" * 70)
    print("CRITICAL ERROR: GEMINI_API_KEY environment variable NOT SET.")
    print("Please ensure your .env file is in the same directory as app.py")
    print("and contains: GEMINI_API_KEY='YOUR_ACTUAL_GEMINI_API_KEY_HERE'")
    print("-" * 70)
    print("\n" * 3)

# --- Define Gemini API URLs ---
GEMINI_TEXT_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
GEMINI_VISION_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# --- Configure Tesseract CMD Path ---
# Read Tesseract path from environment variable, defaulting to common Linux path
# This path should match the installation location within the Dockerfile
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD_PATH', '/usr/bin/tesseract')

@app.route('/upload', methods=['POST', 'OPTIONS'])
@cross_origin()
def upload_file():
    """
    Handles file uploads (PDFs and DOCX for text extraction, images for text extraction via OCR).
    Supports CORS preflight requests (OPTIONS).
    """
    if request.method == 'OPTIONS':
        return '', 200

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'}), 400

    mimetype = file.mimetype
    extracted_text = ""
    file_type_processed = "unknown"
    image_base64 = None # To store base64 for images if needed for multimodal AI
    message = "" # Initialize message for success/failure feedback

    try:
        file_bytes = file.read() # Read file content once
        file_io = io.BytesIO(file_bytes)

        # --- Handle PDF Files ---
        if mimetype == 'application/pdf':
            pdf_reader = PyPDF2.PdfReader(file_io)
            for page in pdf_reader.pages:
                extracted_text_from_page = page.extract_text()
                if extracted_text_from_page:
                    extracted_text += extracted_text_from_page + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from PDF. It might be image-based or empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'pdf'
            message = 'PDF processed successfully!'

        # --- Handle DOCX Files ---
        elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            document = Document(file_io)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from DOCX. It might be empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'docx'
            message = 'DOCX processed successfully!'

        # --- Handle Image Files (Perform OCR and generate base64) ---
        elif mimetype.startswith('image/'):
            try:
                img = Image.open(file_io)
                # Convert image to RGB if it's not (e.g., for PNGs with alpha channel)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Perform OCR
                extracted_text = pytesseract.image_to_string(img)
                
                # Generate base64 for multimodal AI
                buffered = io.BytesIO()
                img.save(buffered, format="PNG") # Save as PNG for consistent base64
                image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                if not extracted_text.strip():
                    message = 'Image uploaded. No readable text found via OCR, but image data is available for multimodal AI.'
                else:
                    message = 'Image processed successfully via OCR!'
                file_type_processed = 'image'

            except pytesseract.TesseractNotFoundError:
                print("Tesseract OCR engine not found. Please ensure Tesseract is installed and its path is correctly configured.")
                return jsonify({'success': False, 'message': 'Tesseract OCR engine not found. Cannot process image for text. Please ensure Tesseract is installed on your server.'}), 500
            except Exception as e:
                print(f"Error during image processing: {e}")
                return jsonify({'success': False, 'message': f'Error during image processing: {str(e)}'}), 500
            
        else:
            message = f'Unsupported file type: {file.filename}. Please upload a PDF, DOCX, or an image.'
            return jsonify({'success': False, 'message': message}), 400

        # Return success response with extracted text, file type, and image base64 if available
        return jsonify({
            'success': True,
            'message': message, # Use the message generated above
            'extractedText': extracted_text,
            'fileType': file_type_processed,
            'imageBase64': image_base64 # Send base64 back for images
        }), 200

    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({'success': False, 'message': f'Server error during file processing: {str(e)}'}), 500

@app.route('/summarize', methods=['POST', 'OPTIONS'])
@cross_origin()
def summarize():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'summary': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # Get image base64 data
        department = data.get('department', '')

        # Check if any content (text or images) is provided
        if not documents and not images_b64:
            return jsonify({'summary': 'No text documents or images provided to summarize.'}), 400

        contents_parts = []
        
        # Build prompt for summarization, explicitly mentioning images if present
        summary_prompt_text = """As UniStudy AI, your purpose is to distill complex information into easily digestible summaries.
Please provide a concise, well-structured, and highly informative summary of the following content.
Your summary should be designed for maximum readability and understanding, using Markdown for:
-   **Clear Headings** (e.g., `## Key Points`, `## Overview`)
-   **Bullet points** or **numbered lists** for main ideas and critical details.
-   **Bold text** (`**important term**`) for emphasizing important terms, definitions, or crucial points.
-   Maintain a professional, intelligent, and helpful tone, similar to an to an advanced AI assistant.
"""
        if documents:
            # Fix: Concatenate string literals and then join documents
            summary_prompt_text += "\n\nHere is the text content to summarize:\n" + '\n\n'.join(documents)
        if images_b64:
            summary_prompt_text += f"\n\nHere are the images to summarize:"

        contents_parts.append({"text": summary_prompt_text})

        # Add image data to contents_parts if available
        for img_data_url in images_b64:
            if ',' in img_data_url:
                try:
                    mime_type = img_data_url.split(';')[0].split(':')[1]
                    base64_data = img_data_url.split(',')[1]
                    contents_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /summarize: {e} for {img_data_url[:50]}...")
                    continue
        
        payload = {
            "contents": [{"role": "user", "parts": contents_parts}], # Wrap contents_parts in a user role
            "generationConfig": {"maxOutputTokens": 500}
        }

        # Determine which Gemini model to use based on presence of images
        use_vision_model = bool(images_b64)
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL

        response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        summary = "Could not generate summary."
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            for part in result['candidates'][0]['content'].get('parts'):
                if 'text' in part:
                    summary = part['text']
                    break
        return jsonify({'summary': summary}), 200
    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for summarize: {req_e}")
        return jsonify({'summary': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for summarize: {e}")
        return jsonify({'summary': f'Error: {str(e)}'}), 500

@app.route('/ask', methods=['POST', 'OPTIONS'])
@cross_origin()
def ask():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'answer': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        question = data.get('question', '')
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # This will now contain base64 strings
        department = data.get('department', '')
        chat_history_raw = data.get('chatHistory', []) # Raw history from frontend

        if not question and not images_b64 and not documents:
            return jsonify({'answer': 'No question, images, or documents provided for AI to process.'}), 400

        contents_parts = []
        
        # Add historical chat for context
        for chat_entry in chat_history_raw:
            if chat_entry.get('sender') == 'user':
                parts = []
                if chat_entry.get('message'):
                    parts.append({"text": chat_entry['message']})
                if chat_entry.get('images'): # If user's past message included images
                    for img_data in chat_entry['images']:
                        # Ensure the image data sent from frontend is already a data URL
                        base64_data = img_data['data'].split(',')[1] if ',' in img_data['data'] else img_data['data']
                        parts.append({"inlineData": {"mimeType": img_data['mimeType'], "data": base64_data}})
                if parts:
                    contents_parts.append({"role": "user", "parts": parts})
            elif chat_entry.get('sender') == 'ai':
                if chat_entry.get('message'):
                    contents_parts.append({"role": "model", "parts": [{"text": chat_entry['message']}]})

        # Add current turn's context and question
        current_user_parts = []
        context_text = "\n\n".join(documents)
        
        # Enhanced prompt for asking questions
        base_prompt = """You are UniStudy AI, an expert, helpful, and highly intelligent study assistant.
        Your goal is to provide comprehensive, accurate, and easy-to-understand answers to student inquiries.
        Please format your response meticulously using Markdown to ensure clarity and readability. Include:
        -   **Prominent Headings** (e.g., `## Main Topic`, `### Sub-topic`) to organize information logically.
        -   **Bullet points** (`- Item`) or **numbered lists** (`1. Item`) for presenting key information, steps, or examples.
        -   **Bold text** (`**important term**`) for significant terms, definitions, or crucial points.
        -   **Code blocks** (```python\nprint("hello")\n```) where appropriate for code, formulas, or structured data.
        -   Adopt a friendly, conversational, yet professional and authoritative tone, akin to a sophisticated AI educator.
        -   If asked a complex question, consider breaking down the answer into logical sections.
        -   If the question can be answered by listing items, use bullet points or numbered lists.
        -   Always strive to provide the most relevant and complete information based on the provided context (documents/images) or your general knowledge if no context is given.
"""

        if context_text:
            # Fix: Concatenate string literals and then context_text and question
            current_user_parts.append({"text": base_prompt + "\n\nHere is the context:\n" + context_text + "\n\nNow, please answer the following question: " + question})
        elif question:
            current_user_parts.append({"text": base_prompt + "\n\nHere is the question: " + question})
        
        # Add images for the current turn if any
        for img_b64_data_url in images_b64:
            if ',' in img_b64_data_url:
                try:
                    mime_type = img_b64_data_url.split(';')[0].split(':')[1]
                    base64_data = img_b64_data_url.split(',')[1] # Extract base64 part
                    current_user_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /ask: {e} for {img_b64_data_url[:50]}...")
                    continue
        
        # Add department context if available and not already part of document context
        if department:
            department_context_added = False
            for part in current_user_parts:
                if "text" in part:
                    if "The question is related to the" not in part["text"]: # Avoid duplicating context
                        part["text"] = f"The question is related to the '{department}' department.\n\n" + part["text"]
                    department_context_added = True
                    break
            if not department_context_added:
                current_user_parts.insert(0, {"text": f"The question is related to the '{department}' department."})


        if not current_user_parts: # If no text or image parts for the current turn
             return jsonify({'answer': 'No valid input provided for the AI to process.'}), 400

        contents_parts.append({"role": "user", "parts": current_user_parts})

        payload = {
            "contents": contents_parts,
            "generationConfig": {"maxOutputTokens": 800}
        }
        
        # --- Debugging: Print the payload before sending ---
        print("\n--- Gemini API Request Payload ---")
        print(json.dumps(payload, indent=2))
        print("----------------------------------\n")

        # Determine which Gemini model to use
        # Use Vision model if there are any images in the current turn or in chat history
        use_vision_model = bool(images_b64) or any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user')
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL
        
        answer = "Sorry, I could not generate an answer."
        try:
            response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
            response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
            result = response.json()
            
            if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                for part in result['candidates'][0]['content'].get('parts'):
                    if 'text' in part:
                        answer = part['text']
                        break
            elif result.get('error'):
                answer = f"AI Error: {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                print(f"Gemini API Error Response: {result['error']}")
            
            return jsonify({'answer': answer}), 200

        except requests.exceptions.RequestException as req_e: 
            print(f"API Request Error for ask: {req_e}")
            # If vision model failed and no images are involved, try text model
            if use_vision_model and not images_b64 and not any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user'):
                print("Vision model failed for a text-only query. Attempting with gemini-2.0-flash.")
                api_url_to_use = GEMINI_TEXT_API_URL # Fallback to text model
                try:
                    response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
                    response.raise_for_status()
                    result = response.json()
                    if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                        for part in result['candidates'][0]['content'].get('parts'):
                            if 'text' in part:
                                answer = part['text']
                                break
                    elif result.get('error'):
                        answer = f"AI Error (fallback): {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                        print(f"Gemini API Error Response (fallback): {result['error']}")
                    return jsonify({'answer': answer}), 200
                except requests.exceptions.RequestException as fallback_req_e:
                    print(f"API Request Error for fallback: {fallback_req_e}")
                    return jsonify({'answer': f'API Error: {str(fallback_req_e)}. Check your API key, model permissions, and request format.'}), 500
            else:
                return jsonify({'answer': f'API Error: {str(req_e)}. Check your API key, model permissions, and request format.'}), 500
    except Exception as e: 
        print(f"General Error for ask: {e}")
        return jsonify({'answer': f'Error: {str(e)}'}), 500

@app.route('/generate_questions', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_questions():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'questions': [], 'message': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        document_text = data.get('documentText', '')
        department = data.get('department', 'general study')

        if not document_text.strip():
            return jsonify({'questions': [], 'message': 'No document text provided for question generation.'}), 400

        # Enhanced prompt for question generation
        prompt = f"""As UniStudy AI, your task is to act as a brilliant question generator, helping students assess their understanding.
        Generate 3-5 concise, insightful, and highly relevant questions that a student might ask based on the following text from the {department} department.
        These questions should encourage critical thinking and comprehension.
        Provide the questions as a JSON array of strings. Example: ["Question 1?", "Question 2?"].
        Ensure the questions are clear, directly related to the provided text, and varied in scope to cover different aspects of the content.
        Text:
        {document_text}
        """
        
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "ARRAY",
                    "items": {"type": "STRING"}
                },
                "maxOutputTokens": 200 # Limit output to keep questions concise
            }
        }

        response = requests.post(f"{GEMINI_TEXT_API_URL}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        
        generated_questions = []
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            try:
                # The response is already JSON, so parse it directly
                json_string = result['candidates'][0]['content']['parts'][0]['text']
                generated_questions = json.loads(json_string)
                if not isinstance(generated_questions, list):
                    generated_questions = [] # Ensure it's a list
            except json.JSONDecodeError as e:
                print(f"JSON decode error in question generation: {e}")
                generated_questions = []
        
        return jsonify({'questions': generated_questions, 'message': 'Questions generated successfully.'}), 200

    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for generate_questions: {req_e}")
        # Added more specific error message for 403
        if req_e.response and req_e.response.status_code == 403:
            return jsonify({'questions': [], 'message': f'API Error: 403 Forbidden. Your API key might not have permissions for this model or API. Please check Google Cloud Console.'}), 500
        else:
            return jsonify({'questions': [], 'message': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for generate_questions: {e}")
        return jsonify({'questions': [], 'message': f'Error: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='localhost', port=5000)
from flask import Flask, request, jsonify
from flask_cors import CORS, cross_origin
import PyPDF2
import io
import requests
import os
from dotenv import load_dotenv
import base64
from docx import Document
from PIL import Image
import pytesseract
import json # Import json for structured output

# --- Load environment variables ---
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# --- Retrieve API Key ---
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    print("\n" * 3)
    print("-" * 70)
    print("CRITICAL ERROR: GEMINI_API_KEY environment variable NOT SET.")
    print("Please ensure your .env file is in the same directory as app.py")
    print("and contains: GEMINI_API_KEY='YOUR_ACTUAL_GEMINI_API_KEY_HERE'")
    print("-" * 70)
    print("\n" * 3)

# --- Define Gemini API URLs ---
GEMINI_TEXT_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
GEMINI_VISION_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# --- Configure Tesseract CMD Path ---
# Read Tesseract path from environment variable, defaulting to common Linux path
# This path should match the installation location within the Dockerfile
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD_PATH', '/usr/bin/tesseract')

@app.route('/upload', methods=['POST', 'OPTIONS'])
@cross_origin()
def upload_file():
    """
    Handles file uploads (PDFs and DOCX for text extraction, images for text extraction via OCR).
    Supports CORS preflight requests (OPTIONS).
    """
    if request.method == 'OPTIONS':
        return '', 200

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'}), 400

    mimetype = file.mimetype
    extracted_text = ""
    file_type_processed = "unknown"
    image_base64 = None # To store base64 for images if needed for multimodal AI
    message = "" # Initialize message for success/failure feedback

    try:
        file_bytes = file.read() # Read file content once
        file_io = io.BytesIO(file_bytes)

        # --- Handle PDF Files ---
        if mimetype == 'application/pdf':
            pdf_reader = PyPDF2.PdfReader(file_io)
            for page in pdf_reader.pages:
                extracted_text_from_page = page.extract_text()
                if extracted_text_from_page:
                    extracted_text += extracted_text_from_page + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from PDF. It might be image-based or empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'pdf'
            message = 'PDF processed successfully!'

        # --- Handle DOCX Files ---
        elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            document = Document(file_io)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from DOCX. It might be empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'docx'
            message = 'DOCX processed successfully!'

        # --- Handle Image Files (Perform OCR and generate base64) ---
        elif mimetype.startswith('image/'):
            try:
                img = Image.open(file_io)
                # Convert image to RGB if it's not (e.g., for PNGs with alpha channel)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Perform OCR
                extracted_text = pytesseract.image_to_string(img)
                
                # Generate base64 for multimodal AI
                buffered = io.BytesIO()
                img.save(buffered, format="PNG") # Save as PNG for consistent base64
                image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                if not extracted_text.strip():
                    message = 'Image uploaded. No readable text found via OCR, but image data is available for multimodal AI.'
                else:
                    message = 'Image processed successfully via OCR!'
                file_type_processed = 'image'

            except pytesseract.TesseractNotFoundError:
                print("Tesseract OCR engine not found. Please ensure Tesseract is installed and its path is correctly configured.")
                return jsonify({'success': False, 'message': 'Tesseract OCR engine not found. Cannot process image for text. Please ensure Tesseract is installed on your server.'}), 500
            except Exception as e:
                print(f"Error during image processing: {e}")
                return jsonify({'success': False, 'message': f'Error during image processing: {str(e)}'}), 500
            
        else:
            message = f'Unsupported file type: {file.filename}. Please upload a PDF, DOCX, or an image.'
            return jsonify({'success': False, 'message': message}), 400

        # Return success response with extracted text, file type, and image base64 if available
        return jsonify({
            'success': True,
            'message': message, # Use the message generated above
            'extractedText': extracted_text,
            'fileType': file_type_processed,
            'imageBase64': image_base64 # Send base64 back for images
        }), 200

    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({'success': False, 'message': f'Server error during file processing: {str(e)}'}), 500

@app.route('/summarize', methods=['POST', 'OPTIONS'])
@cross_origin()
def summarize():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'summary': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # Get image base64 data
        department = data.get('department', '')

        # Check if any content (text or images) is provided
        if not documents and not images_b64:
            return jsonify({'summary': 'No text documents or images provided to summarize.'}), 400

        contents_parts = []
        
        # Build prompt for summarization, explicitly mentioning images if present
        summary_prompt_text = """As UniStudy AI, your purpose is to distill complex information into easily digestible summaries.
Please provide a concise, well-structured, and highly informative summary of the following content.
Your summary should be designed for maximum readability and understanding, using Markdown for:
-   **Clear Headings** (e.g., `## Key Points`, `## Overview`)
-   **Bullet points** or **numbered lists** for main ideas and critical details.
-   **Bold text** (`**important term**`) for emphasizing important terms, definitions, or crucial points.
-   Maintain a professional, intelligent, and helpful tone, similar to an to an advanced AI assistant.
"""
        if documents:
            # Fix: Concatenate string literals and then join documents
            summary_prompt_text += "\n\nHere is the text content to summarize:\n" + '\n\n'.join(documents)
        if images_b64:
            summary_prompt_text += f"\n\nHere are the images to summarize:"

        contents_parts.append({"text": summary_prompt_text})

        # Add image data to contents_parts if available
        for img_data_url in images_b64:
            if ',' in img_data_url:
                try:
                    mime_type = img_data_url.split(';')[0].split(':')[1]
                    base64_data = img_data_url.split(',')[1]
                    contents_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /summarize: {e} for {img_data_url[:50]}...")
                    continue
        
        payload = {
            "contents": [{"role": "user", "parts": contents_parts}], # Wrap contents_parts in a user role
            "generationConfig": {"maxOutputTokens": 500}
        }

        # Determine which Gemini model to use based on presence of images
        use_vision_model = bool(images_b64)
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL

        response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        summary = "Could not generate summary."
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            for part in result['candidates'][0]['content'].get('parts'):
                if 'text' in part:
                    summary = part['text']
                    break
        return jsonify({'summary': summary}), 200
    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for summarize: {req_e}")
        return jsonify({'summary': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for summarize: {e}")
        return jsonify({'summary': f'Error: {str(e)}'}), 500

@app.route('/ask', methods=['POST', 'OPTIONS'])
@cross_origin()
def ask():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'answer': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        question = data.get('question', '')
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # This will now contain base64 strings
        department = data.get('department', '')
        chat_history_raw = data.get('chatHistory', []) # Raw history from frontend

        if not question and not images_b64 and not documents:
            return jsonify({'answer': 'No question, images, or documents provided for AI to process.'}), 400

        contents_parts = []
        
        # Add historical chat for context
        for chat_entry in chat_history_raw:
            if chat_entry.get('sender') == 'user':
                parts = []
                if chat_entry.get('message'):
                    parts.append({"text": chat_entry['message']})
                if chat_entry.get('images'): # If user's past message included images
                    for img_data in chat_entry['images']:
                        # Ensure the image data sent from frontend is already a data URL
                        base64_data = img_data['data'].split(',')[1] if ',' in img_data['data'] else img_data['data']
                        parts.append({"inlineData": {"mimeType": img_data['mimeType'], "data": base64_data}})
                if parts:
                    contents_parts.append({"role": "user", "parts": parts})
            elif chat_entry.get('sender') == 'ai':
                if chat_entry.get('message'):
                    contents_parts.append({"role": "model", "parts": [{"text": chat_entry['message']}]})

        # Add current turn's context and question
        current_user_parts = []
        context_text = "\n\n".join(documents)
        
        # Enhanced prompt for asking questions
        base_prompt = """You are UniStudy AI, an expert, helpful, and highly intelligent study assistant.
        Your goal is to provide comprehensive, accurate, and easy-to-understand answers to student inquiries.
        Please format your response meticulously using Markdown to ensure clarity and readability. Include:
        -   **Prominent Headings** (e.g., `## Main Topic`, `### Sub-topic`) to organize information logically.
        -   **Bullet points** (`- Item`) or **numbered lists** (`1. Item`) for presenting key information, steps, or examples.
        -   **Bold text** (`**important term**`) for significant terms, definitions, or crucial points.
        -   **Code blocks** (```python\nprint("hello")\n```) where appropriate for code, formulas, or structured data.
        -   Adopt a friendly, conversational, yet professional and authoritative tone, akin to a sophisticated AI educator.
        -   If asked a complex question, consider breaking down the answer into logical sections.
        -   If the question can be answered by listing items, use bullet points or numbered lists.
        -   Always strive to provide the most relevant and complete information based on the provided context (documents/images) or your general knowledge if no context is given.
"""

        if context_text:
            # Fix: Concatenate string literals and then context_text and question
            current_user_parts.append({"text": base_prompt + "\n\nHere is the context:\n" + context_text + "\n\nNow, please answer the following question: " + question})
        elif question:
            current_user_parts.append({"text": base_prompt + "\n\nHere is the question: " + question})
        
        # Add images for the current turn if any
        for img_b64_data_url in images_b64:
            if ',' in img_b64_data_url:
                try:
                    mime_type = img_b64_data_url.split(';')[0].split(':')[1]
                    base64_data = img_b64_data_url.split(',')[1] # Extract base64 part
                    current_user_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /ask: {e} for {img_b64_data_url[:50]}...")
                    continue
        
        # Add department context if available and not already part of document context
        if department:
            department_context_added = False
            for part in current_user_parts:
                if "text" in part:
                    if "The question is related to the" not in part["text"]: # Avoid duplicating context
                        part["text"] = f"The question is related to the '{department}' department.\n\n" + part["text"]
                    department_context_added = True
                    break
            if not department_context_added:
                current_user_parts.insert(0, {"text": f"The question is related to the '{department}' department."})


        if not current_user_parts: # If no text or image parts for the current turn
             return jsonify({'answer': 'No valid input provided for the AI to process.'}), 400

        contents_parts.append({"role": "user", "parts": current_user_parts})

        payload = {
            "contents": contents_parts,
            "generationConfig": {"maxOutputTokens": 800}
        }
        
        # --- Debugging: Print the payload before sending ---
        print("\n--- Gemini API Request Payload ---")
        print(json.dumps(payload, indent=2))
        print("----------------------------------\n")

        # Determine which Gemini model to use
        # Use Vision model if there are any images in the current turn or in chat history
        use_vision_model = bool(images_b64) or any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user')
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL
        
        answer = "Sorry, I could not generate an answer."
        try:
            response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
            response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
            result = response.json()
            
            if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                for part in result['candidates'][0]['content'].get('parts'):
                    if 'text' in part:
                        answer = part['text']
                        break
            elif result.get('error'):
                answer = f"AI Error: {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                print(f"Gemini API Error Response: {result['error']}")
            
            return jsonify({'answer': answer}), 200

        except requests.exceptions.RequestException as req_e: 
            print(f"API Request Error for ask: {req_e}")
            # If vision model failed and no images are involved, try text model
            if use_vision_model and not images_b64 and not any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user'):
                print("Vision model failed for a text-only query. Attempting with gemini-2.0-flash.")
                api_url_to_use = GEMINI_TEXT_API_URL # Fallback to text model
                try:
                    response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
                    response.raise_for_status()
                    result = response.json()
                    if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                        for part in result['candidates'][0]['content'].get('parts'):
                            if 'text' in part:
                                answer = part['text']
                                break
                    elif result.get('error'):
                        answer = f"AI Error (fallback): {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                        print(f"Gemini API Error Response (fallback): {result['error']}")
                    return jsonify({'answer': answer}), 200
                except requests.exceptions.RequestException as fallback_req_e:
                    print(f"API Request Error for fallback: {fallback_req_e}")
                    return jsonify({'answer': f'API Error: {str(fallback_req_e)}. Check your API key, model permissions, and request format.'}), 500
            else:
                return jsonify({'answer': f'API Error: {str(req_e)}. Check your API key, model permissions, and request format.'}), 500
    except Exception as e: 
        print(f"General Error for ask: {e}")
        return jsonify({'answer': f'Error: {str(e)}'}), 500

@app.route('/generate_questions', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_questions():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'questions': [], 'message': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        document_text = data.get('documentText', '')
        department = data.get('department', 'general study')

        if not document_text.strip():
            return jsonify({'questions': [], 'message': 'No document text provided for question generation.'}), 400

        # Enhanced prompt for question generation
        prompt = f"""As UniStudy AI, your task is to act as a brilliant question generator, helping students assess their understanding.
        Generate 3-5 concise, insightful, and highly relevant questions that a student might ask based on the following text from the {department} department.
        These questions should encourage critical thinking and comprehension.
        Provide the questions as a JSON array of strings. Example: ["Question 1?", "Question 2?"].
        Ensure the questions are clear, directly related to the provided text, and varied in scope to cover different aspects of the content.
        Text:
        {document_text}
        """
        
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "ARRAY",
                    "items": {"type": "STRING"}
                },
                "maxOutputTokens": 200 # Limit output to keep questions concise
            }
        }

        response = requests.post(f"{GEMINI_TEXT_API_URL}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        
        generated_questions = []
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            try:
                # The response is already JSON, so parse it directly
                json_string = result['candidates'][0]['content']['parts'][0]['text']
                generated_questions = json.loads(json_string)
                if not isinstance(generated_questions, list):
                    generated_questions = [] # Ensure it's a list
            except json.JSONDecodeError as e:
                print(f"JSON decode error in question generation: {e}")
                generated_questions = []
        
        return jsonify({'questions': generated_questions, 'message': 'Questions generated successfully.'}), 200

    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for generate_questions: {req_e}")
        # Added more specific error message for 403
        if req_e.response and req_e.response.status_code == 403:
            return jsonify({'questions': [], 'message': f'API Error: 403 Forbidden. Your API key might not have permissions for this model or API. Please check Google Cloud Console.'}), 500
        else:
            return jsonify({'questions': [], 'message': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for generate_questions: {e}")
        return jsonify({'questions': [], 'message': f'Error: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='localhost', port=5000)
from flask import Flask, request, jsonify
from flask_cors import CORS, cross_origin
import PyPDF2
import io
import requests
import os
from dotenv import load_dotenv
import base64
from docx import Document
from PIL import Image
import pytesseract
import json # Import json for structured output

# --- Load environment variables ---
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# --- Retrieve API Key ---
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    print("\n" * 3)
    print("-" * 70)
    print("CRITICAL ERROR: GEMINI_API_KEY environment variable NOT SET.")
    print("Please ensure your .env file is in the same directory as app.py")
    print("and contains: GEMINI_API_KEY='YOUR_ACTUAL_GEMINI_API_KEY_HERE'")
    print("-" * 70)
    print("\n" * 3)

# --- Define Gemini API URLs ---
GEMINI_TEXT_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
GEMINI_VISION_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# --- Configure Tesseract CMD Path ---
# Read Tesseract path from environment variable, defaulting to common Linux path
# This path should match the installation location within the Dockerfile
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD_PATH', '/usr/bin/tesseract')

@app.route('/upload', methods=['POST', 'OPTIONS'])
@cross_origin()
def upload_file():
    """
    Handles file uploads (PDFs and DOCX for text extraction, images for text extraction via OCR).
    Supports CORS preflight requests (OPTIONS).
    """
    if request.method == 'OPTIONS':
        return '', 200

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'}), 400

    mimetype = file.mimetype
    extracted_text = ""
    file_type_processed = "unknown"
    image_base64 = None # To store base64 for images if needed for multimodal AI
    message = "" # Initialize message for success/failure feedback

    try:
        file_bytes = file.read() # Read file content once
        file_io = io.BytesIO(file_bytes)

        # --- Handle PDF Files ---
        if mimetype == 'application/pdf':
            pdf_reader = PyPDF2.PdfReader(file_io)
            for page in pdf_reader.pages:
                extracted_text_from_page = page.extract_text()
                if extracted_text_from_page:
                    extracted_text += extracted_text_from_page + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from PDF. It might be image-based or empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'pdf'
            message = 'PDF processed successfully!'

        # --- Handle DOCX Files ---
        elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            document = Document(file_io)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from DOCX. It might be empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'docx'
            message = 'DOCX processed successfully!'

        # --- Handle Image Files (Perform OCR and generate base64) ---
        elif mimetype.startswith('image/'):
            try:
                img = Image.open(file_io)
                # Convert image to RGB if it's not (e.g., for PNGs with alpha channel)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Perform OCR
                extracted_text = pytesseract.image_to_string(img)
                
                # Generate base64 for multimodal AI
                buffered = io.BytesIO()
                img.save(buffered, format="PNG") # Save as PNG for consistent base64
                image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                if not extracted_text.strip():
                    message = 'Image uploaded. No readable text found via OCR, but image data is available for multimodal AI.'
                else:
                    message = 'Image processed successfully via OCR!'
                file_type_processed = 'image'

            except pytesseract.TesseractNotFoundError:
                print("Tesseract OCR engine not found. Please ensure Tesseract is installed and its path is correctly configured.")
                return jsonify({'success': False, 'message': 'Tesseract OCR engine not found. Cannot process image for text. Please ensure Tesseract is installed on your server.'}), 500
            except Exception as e:
                print(f"Error during image processing: {e}")
                return jsonify({'success': False, 'message': f'Error during image processing: {str(e)}'}), 500
            
        else:
            message = f'Unsupported file type: {file.filename}. Please upload a PDF, DOCX, or an image.'
            return jsonify({'success': False, 'message': message}), 400

        # Return success response with extracted text, file type, and image base64 if available
        return jsonify({
            'success': True,
            'message': message, # Use the message generated above
            'extractedText': extracted_text,
            'fileType': file_type_processed,
            'imageBase64': image_base64 # Send base64 back for images
        }), 200

    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({'success': False, 'message': f'Server error during file processing: {str(e)}'}), 500

@app.route('/summarize', methods=['POST', 'OPTIONS'])
@cross_origin()
def summarize():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'summary': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # Get image base64 data
        department = data.get('department', '')

        # Check if any content (text or images) is provided
        if not documents and not images_b64:
            return jsonify({'summary': 'No text documents or images provided to summarize.'}), 400

        contents_parts = []
        
        # Build prompt for summarization, explicitly mentioning images if present
        summary_prompt_text = """As UniStudy AI, your purpose is to distill complex information into easily digestible summaries.
Please provide a concise, well-structured, and highly informative summary of the following content.
Your summary should be designed for maximum readability and understanding, using Markdown for:
-   **Clear Headings** (e.g., `## Key Points`, `## Overview`)
-   **Bullet points** or **numbered lists** for main ideas and critical details.
-   **Bold text** (`**important term**`) for emphasizing important terms, definitions, or crucial points.
-   Maintain a professional, intelligent, and helpful tone, similar to an to an advanced AI assistant.
"""
        if documents:
            # Fix: Concatenate string literals and then join documents
            summary_prompt_text += "\n\nHere is the text content to summarize:\n" + '\n\n'.join(documents)
        if images_b64:
            summary_prompt_text += f"\n\nHere are the images to summarize:"

        contents_parts.append({"text": summary_prompt_text})

        # Add image data to contents_parts if available
        for img_data_url in images_b64:
            if ',' in img_data_url:
                try:
                    mime_type = img_data_url.split(';')[0].split(':')[1]
                    base64_data = img_data_url.split(',')[1]
                    contents_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /summarize: {e} for {img_data_url[:50]}...")
                    continue
        
        payload = {
            "contents": [{"role": "user", "parts": contents_parts}], # Wrap contents_parts in a user role
            "generationConfig": {"maxOutputTokens": 500}
        }

        # Determine which Gemini model to use based on presence of images
        use_vision_model = bool(images_b64)
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL

        response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        summary = "Could not generate summary."
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            for part in result['candidates'][0]['content'].get('parts'):
                if 'text' in part:
                    summary = part['text']
                    break
        return jsonify({'summary': summary}), 200
    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for summarize: {req_e}")
        return jsonify({'summary': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for summarize: {e}")
        return jsonify({'summary': f'Error: {str(e)}'}), 500

@app.route('/ask', methods=['POST', 'OPTIONS'])
@cross_origin()
def ask():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'answer': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        question = data.get('question', '')
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # This will now contain base64 strings
        department = data.get('department', '')
        chat_history_raw = data.get('chatHistory', []) # Raw history from frontend

        if not question and not images_b64 and not documents:
            return jsonify({'answer': 'No question, images, or documents provided for AI to process.'}), 400

        contents_parts = []
        
        # Add historical chat for context
        for chat_entry in chat_history_raw:
            if chat_entry.get('sender') == 'user':
                parts = []
                if chat_entry.get('message'):
                    parts.append({"text": chat_entry['message']})
                if chat_entry.get('images'): # If user's past message included images
                    for img_data in chat_entry['images']:
                        # Ensure the image data sent from frontend is already a data URL
                        base64_data = img_data['data'].split(',')[1] if ',' in img_data['data'] else img_data['data']
                        parts.append({"inlineData": {"mimeType": img_data['mimeType'], "data": base64_data}})
                if parts:
                    contents_parts.append({"role": "user", "parts": parts})
            elif chat_entry.get('sender') == 'ai':
                if chat_entry.get('message'):
                    contents_parts.append({"role": "model", "parts": [{"text": chat_entry['message']}]})

        # Add current turn's context and question
        current_user_parts = []
        context_text = "\n\n".join(documents)
        
        # Enhanced prompt for asking questions
        base_prompt = """You are UniStudy AI, an expert, helpful, and highly intelligent study assistant.
        Your goal is to provide comprehensive, accurate, and easy-to-understand answers to student inquiries.
        Please format your response meticulously using Markdown to ensure clarity and readability. Include:
        -   **Prominent Headings** (e.g., `## Main Topic`, `### Sub-topic`) to organize information logically.
        -   **Bullet points** (`- Item`) or **numbered lists** (`1. Item`) for presenting key information, steps, or examples.
        -   **Bold text** (`**important term**`) for significant terms, definitions, or crucial points.
        -   **Code blocks** (```python\nprint("hello")\n```) where appropriate for code, formulas, or structured data.
        -   Adopt a friendly, conversational, yet professional and authoritative tone, akin to a sophisticated AI educator.
        -   If asked a complex question, consider breaking down the answer into logical sections.
        -   If the question can be answered by listing items, use bullet points or numbered lists.
        -   Always strive to provide the most relevant and complete information based on the provided context (documents/images) or your general knowledge if no context is given.
"""

        if context_text:
            # Fix: Concatenate string literals and then context_text and question
            current_user_parts.append({"text": base_prompt + "\n\nHere is the context:\n" + context_text + "\n\nNow, please answer the following question: " + question})
        elif question:
            current_user_parts.append({"text": base_prompt + "\n\nHere is the question: " + question})
        
        # Add images for the current turn if any
        for img_b64_data_url in images_b64:
            if ',' in img_b64_data_url:
                try:
                    mime_type = img_b64_data_url.split(';')[0].split(':')[1]
                    base64_data = img_b64_data_url.split(',')[1] # Extract base64 part
                    current_user_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /ask: {e} for {img_b64_data_url[:50]}...")
                    continue
        
        # Add department context if available and not already part of document context
        if department:
            department_context_added = False
            for part in current_user_parts:
                if "text" in part:
                    if "The question is related to the" not in part["text"]: # Avoid duplicating context
                        part["text"] = f"The question is related to the '{department}' department.\n\n" + part["text"]
                    department_context_added = True
                    break
            if not department_context_added:
                current_user_parts.insert(0, {"text": f"The question is related to the '{department}' department."})


        if not current_user_parts: # If no text or image parts for the current turn
             return jsonify({'answer': 'No valid input provided for the AI to process.'}), 400

        contents_parts.append({"role": "user", "parts": current_user_parts})

        payload = {
            "contents": contents_parts,
            "generationConfig": {"maxOutputTokens": 800}
        }
        
        # --- Debugging: Print the payload before sending ---
        print("\n--- Gemini API Request Payload ---")
        print(json.dumps(payload, indent=2))
        print("----------------------------------\n")

        # Determine which Gemini model to use
        # Use Vision model if there are any images in the current turn or in chat history
        use_vision_model = bool(images_b64) or any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user')
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL
        
        answer = "Sorry, I could not generate an answer."
        try:
            response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
            response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
            result = response.json()
            
            if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                for part in result['candidates'][0]['content'].get('parts'):
                    if 'text' in part:
                        answer = part['text']
                        break
            elif result.get('error'):
                answer = f"AI Error: {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                print(f"Gemini API Error Response: {result['error']}")
            
            return jsonify({'answer': answer}), 200

        except requests.exceptions.RequestException as req_e: 
            print(f"API Request Error for ask: {req_e}")
            # If vision model failed and no images are involved, try text model
            if use_vision_model and not images_b64 and not any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user'):
                print("Vision model failed for a text-only query. Attempting with gemini-2.0-flash.")
                api_url_to_use = GEMINI_TEXT_API_URL # Fallback to text model
                try:
                    response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
                    response.raise_for_status()
                    result = response.json()
                    if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                        for part in result['candidates'][0]['content'].get('parts'):
                            if 'text' in part:
                                answer = part['text']
                                break
                    elif result.get('error'):
                        answer = f"AI Error (fallback): {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                        print(f"Gemini API Error Response (fallback): {result['error']}")
                    return jsonify({'answer': answer}), 200
                except requests.exceptions.RequestException as fallback_req_e:
                    print(f"API Request Error for fallback: {fallback_req_e}")
                    return jsonify({'answer': f'API Error: {str(fallback_req_e)}. Check your API key, model permissions, and request format.'}), 500
            else:
                return jsonify({'answer': f'API Error: {str(req_e)}. Check your API key, model permissions, and request format.'}), 500
    except Exception as e: 
        print(f"General Error for ask: {e}")
        return jsonify({'answer': f'Error: {str(e)}'}), 500

@app.route('/generate_questions', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_questions():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'questions': [], 'message': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        document_text = data.get('documentText', '')
        department = data.get('department', 'general study')

        if not document_text.strip():
            return jsonify({'questions': [], 'message': 'No document text provided for question generation.'}), 400

        # Enhanced prompt for question generation
        prompt = f"""As UniStudy AI, your task is to act as a brilliant question generator, helping students assess their understanding.
        Generate 3-5 concise, insightful, and highly relevant questions that a student might ask based on the following text from the {department} department.
        These questions should encourage critical thinking and comprehension.
        Provide the questions as a JSON array of strings. Example: ["Question 1?", "Question 2?"].
        Ensure the questions are clear, directly related to the provided text, and varied in scope to cover different aspects of the content.
        Text:
        {document_text}
        """
        
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "ARRAY",
                    "items": {"type": "STRING"}
                },
                "maxOutputTokens": 200 # Limit output to keep questions concise
            }
        }

        response = requests.post(f"{GEMINI_TEXT_API_URL}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        
        generated_questions = []
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            try:
                # The response is already JSON, so parse it directly
                json_string = result['candidates'][0]['content']['parts'][0]['text']
                generated_questions = json.loads(json_string)
                if not isinstance(generated_questions, list):
                    generated_questions = [] # Ensure it's a list
            except json.JSONDecodeError as e:
                print(f"JSON decode error in question generation: {e}")
                generated_questions = []
        
        return jsonify({'questions': generated_questions, 'message': 'Questions generated successfully.'}), 200

    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for generate_questions: {req_e}")
        # Added more specific error message for 403
        if req_e.response and req_e.response.status_code == 403:
            return jsonify({'questions': [], 'message': f'API Error: 403 Forbidden. Your API key might not have permissions for this model or API. Please check Google Cloud Console.'}), 500
        else:
            return jsonify({'questions': [], 'message': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for generate_questions: {e}")
        return jsonify({'questions': [], 'message': f'Error: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='localhost', port=5000)
from flask import Flask, request, jsonify
from flask_cors import CORS, cross_origin
import PyPDF2
import io
import requests
import os
from dotenv import load_dotenv
import base64
from docx import Document
from PIL import Image
import pytesseract
import json # Import json for structured output

# --- Load environment variables ---
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# --- Retrieve API Key ---
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    print("\n" * 3)
    print("-" * 70)
    print("CRITICAL ERROR: GEMINI_API_KEY environment variable NOT SET.")
    print("Please ensure your .env file is in the same directory as app.py")
    print("and contains: GEMINI_API_KEY='YOUR_ACTUAL_GEMINI_API_KEY_HERE'")
    print("-" * 70)
    print("\n" * 3)

# --- Define Gemini API URLs ---
GEMINI_TEXT_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
GEMINI_VISION_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# --- Configure Tesseract CMD Path ---
# Read Tesseract path from environment variable, defaulting to common Linux path
# This path should match the installation location within the Dockerfile
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD_PATH', '/usr/bin/tesseract')

@app.route('/upload', methods=['POST', 'OPTIONS'])
@cross_origin()
def upload_file():
    """
    Handles file uploads (PDFs and DOCX for text extraction, images for text extraction via OCR).
    Supports CORS preflight requests (OPTIONS).
    """
    if request.method == 'OPTIONS':
        return '', 200

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'}), 400

    mimetype = file.mimetype
    extracted_text = ""
    file_type_processed = "unknown"
    image_base64 = None # To store base64 for images if needed for multimodal AI
    message = "" # Initialize message for success/failure feedback

    try:
        file_bytes = file.read() # Read file content once
        file_io = io.BytesIO(file_bytes)

        # --- Handle PDF Files ---
        if mimetype == 'application/pdf':
            pdf_reader = PyPDF2.PdfReader(file_io)
            for page in pdf_reader.pages:
                extracted_text_from_page = page.extract_text()
                if extracted_text_from_page:
                    extracted_text += extracted_text_from_page + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from PDF. It might be image-based or empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'pdf'
            message = 'PDF processed successfully!'

        # --- Handle DOCX Files ---
        elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            document = Document(file_io)
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            if not extracted_text.strip():
                message = 'Could not extract text from DOCX. It might be empty.'
                return jsonify({'success': False, 'message': message}), 400
            file_type_processed = 'docx'
            message = 'DOCX processed successfully!'

        # --- Handle Image Files (Perform OCR and generate base64) ---
        elif mimetype.startswith('image/'):
            try:
                img = Image.open(file_io)
                # Convert image to RGB if it's not (e.g., for PNGs with alpha channel)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Perform OCR
                extracted_text = pytesseract.image_to_string(img)
                
                # Generate base64 for multimodal AI
                buffered = io.BytesIO()
                img.save(buffered, format="PNG") # Save as PNG for consistent base64
                image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                if not extracted_text.strip():
                    message = 'Image uploaded. No readable text found via OCR, but image data is available for multimodal AI.'
                else:
                    message = 'Image processed successfully via OCR!'
                file_type_processed = 'image'

            except pytesseract.TesseractNotFoundError:
                print("Tesseract OCR engine not found. Please ensure Tesseract is installed and its path is correctly configured.")
                return jsonify({'success': False, 'message': 'Tesseract OCR engine not found. Cannot process image for text. Please ensure Tesseract is installed on your server.'}), 500
            except Exception as e:
                print(f"Error during image processing: {e}")
                return jsonify({'success': False, 'message': f'Error during image processing: {str(e)}'}), 500
            
        else:
            message = f'Unsupported file type: {file.filename}. Please upload a PDF, DOCX, or an image.'
            return jsonify({'success': False, 'message': message}), 400

        # Return success response with extracted text, file type, and image base64 if available
        return jsonify({
            'success': True,
            'message': message, # Use the message generated above
            'extractedText': extracted_text,
            'fileType': file_type_processed,
            'imageBase64': image_base64 # Send base64 back for images
        }), 200

    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({'success': False, 'message': f'Server error during file processing: {str(e)}'}), 500

@app.route('/summarize', methods=['POST', 'OPTIONS'])
@cross_origin()
def summarize():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'summary': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # Get image base64 data
        department = data.get('department', '')

        # Check if any content (text or images) is provided
        if not documents and not images_b64:
            return jsonify({'summary': 'No text documents or images provided to summarize.'}), 400

        contents_parts = []
        
        # Build prompt for summarization, explicitly mentioning images if present
        summary_prompt_text = """As UniStudy AI, your purpose is to distill complex information into easily digestible summaries.
Please provide a concise, well-structured, and highly informative summary of the following content.
Your summary should be designed for maximum readability and understanding, using Markdown for:
-   **Clear Headings** (e.g., `## Key Points`, `## Overview`)
-   **Bullet points** or **numbered lists** for main ideas and critical details.
-   **Bold text** (`**important term**`) for emphasizing important terms, definitions, or crucial points.
-   Maintain a professional, intelligent, and helpful tone, similar to an to an advanced AI assistant.
"""
        if documents:
            # Fix: Concatenate string literals and then join documents
            summary_prompt_text += "\n\nHere is the text content to summarize:\n" + '\n\n'.join(documents)
        if images_b64:
            summary_prompt_text += f"\n\nHere are the images to summarize:"

        contents_parts.append({"text": summary_prompt_text})

        # Add image data to contents_parts if available
        for img_data_url in images_b64:
            if ',' in img_data_url:
                try:
                    mime_type = img_data_url.split(';')[0].split(':')[1]
                    base64_data = img_data_url.split(',')[1]
                    contents_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /summarize: {e} for {img_data_url[:50]}...")
                    continue
        
        payload = {
            "contents": [{"role": "user", "parts": contents_parts}], # Wrap contents_parts in a user role
            "generationConfig": {"maxOutputTokens": 500}
        }

        # Determine which Gemini model to use based on presence of images
        use_vision_model = bool(images_b64)
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL

        response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        summary = "Could not generate summary."
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            for part in result['candidates'][0]['content'].get('parts'):
                if 'text' in part:
                    summary = part['text']
                    break
        return jsonify({'summary': summary}), 200
    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for summarize: {req_e}")
        return jsonify({'summary': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for summarize: {e}")
        return jsonify({'summary': f'Error: {str(e)}'}), 500

@app.route('/ask', methods=['POST', 'OPTIONS'])
@cross_origin()
def ask():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'answer': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        question = data.get('question', '')
        documents = data.get('documents', [])
        images_b64 = data.get('images', []) # This will now contain base64 strings
        department = data.get('department', '')
        chat_history_raw = data.get('chatHistory', []) # Raw history from frontend

        if not question and not images_b64 and not documents:
            return jsonify({'answer': 'No question, images, or documents provided for AI to process.'}), 400

        contents_parts = []
        
        # Add historical chat for context
        for chat_entry in chat_history_raw:
            if chat_entry.get('sender') == 'user':
                parts = []
                if chat_entry.get('message'):
                    parts.append({"text": chat_entry['message']})
                if chat_entry.get('images'): # If user's past message included images
                    for img_data in chat_entry['images']:
                        # Ensure the image data sent from frontend is already a data URL
                        base64_data = img_data['data'].split(',')[1] if ',' in img_data['data'] else img_data['data']
                        parts.append({"inlineData": {"mimeType": img_data['mimeType'], "data": base64_data}})
                if parts:
                    contents_parts.append({"role": "user", "parts": parts})
            elif chat_entry.get('sender') == 'ai':
                if chat_entry.get('message'):
                    contents_parts.append({"role": "model", "parts": [{"text": chat_entry['message']}]})

        # Add current turn's context and question
        current_user_parts = []
        context_text = "\n\n".join(documents)
        
        # Enhanced prompt for asking questions
        base_prompt = """You are UniStudy AI, an expert, helpful, and highly intelligent study assistant.
        Your goal is to provide comprehensive, accurate, and easy-to-understand answers to student inquiries.
        Please format your response meticulously using Markdown to ensure clarity and readability. Include:
        -   **Prominent Headings** (e.g., `## Main Topic`, `### Sub-topic`) to organize information logically.
        -   **Bullet points** (`- Item`) or **numbered lists** (`1. Item`) for presenting key information, steps, or examples.
        -   **Bold text** (`**important term**`) for significant terms, definitions, or crucial points.
        -   **Code blocks** (```python\nprint("hello")\n```) where appropriate for code, formulas, or structured data.
        -   Adopt a friendly, conversational, yet professional and authoritative tone, akin to a sophisticated AI educator.
        -   If asked a complex question, consider breaking down the answer into logical sections.
        -   If the question can be answered by listing items, use bullet points or numbered lists.
        -   Always strive to provide the most relevant and complete information based on the provided context (documents/images) or your general knowledge if no context is given.
"""

        if context_text:
            # Fix: Concatenate string literals and then context_text and question
            current_user_parts.append({"text": base_prompt + "\n\nHere is the context:\n" + context_text + "\n\nNow, please answer the following question: " + question})
        elif question:
            current_user_parts.append({"text": base_prompt + "\n\nHere is the question: " + question})
        
        # Add images for the current turn if any
        for img_b64_data_url in images_b64:
            if ',' in img_b64_data_url:
                try:
                    mime_type = img_b64_data_url.split(';')[0].split(':')[1]
                    base64_data = img_b64_data_url.split(',')[1] # Extract base64 part
                    current_user_parts.append({
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": base64_data
                        }
                    })
                except Exception as e:
                    print(f"Error processing image data URL in /ask: {e} for {img_b64_data_url[:50]}...")
                    continue
        
        # Add department context if available and not already part of document context
        if department:
            department_context_added = False
            for part in current_user_parts:
                if "text" in part:
                    if "The question is related to the" not in part["text"]: # Avoid duplicating context
                        part["text"] = f"The question is related to the '{department}' department.\n\n" + part["text"]
                    department_context_added = True
                    break
            if not department_context_added:
                current_user_parts.insert(0, {"text": f"The question is related to the '{department}' department."})


        if not current_user_parts: # If no text or image parts for the current turn
             return jsonify({'answer': 'No valid input provided for the AI to process.'}), 400

        contents_parts.append({"role": "user", "parts": current_user_parts})

        payload = {
            "contents": contents_parts,
            "generationConfig": {"maxOutputTokens": 800}
        }
        
        # --- Debugging: Print the payload before sending ---
        print("\n--- Gemini API Request Payload ---")
        print(json.dumps(payload, indent=2))
        print("----------------------------------\n")

        # Determine which Gemini model to use
        # Use Vision model if there are any images in the current turn or in chat history
        use_vision_model = bool(images_b64) or any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user')
        api_url_to_use = GEMINI_VISION_API_URL if use_vision_model else GEMINI_TEXT_API_URL
        
        answer = "Sorry, I could not generate an answer."
        try:
            response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
            response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
            result = response.json()
            
            if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                for part in result['candidates'][0]['content'].get('parts'):
                    if 'text' in part:
                        answer = part['text']
                        break
            elif result.get('error'):
                answer = f"AI Error: {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                print(f"Gemini API Error Response: {result['error']}")
            
            return jsonify({'answer': answer}), 200

        except requests.exceptions.RequestException as req_e: 
            print(f"API Request Error for ask: {req_e}")
            # If vision model failed and no images are involved, try text model
            if use_vision_model and not images_b64 and not any('images' in entry for entry in chat_history_raw if entry.get('sender') == 'user'):
                print("Vision model failed for a text-only query. Attempting with gemini-2.0-flash.")
                api_url_to_use = GEMINI_TEXT_API_URL # Fallback to text model
                try:
                    response = requests.post(f"{api_url_to_use}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
                    response.raise_for_status()
                    result = response.json()
                    if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
                        for part in result['candidates'][0]['content'].get('parts'):
                            if 'text' in part:
                                answer = part['text']
                                break
                    elif result.get('error'):
                        answer = f"AI Error (fallback): {result['error'].get('message', 'Unknown API error')}. Code: {result['error'].get('code', 'N/A')}"
                        print(f"Gemini API Error Response (fallback): {result['error']}")
                    return jsonify({'answer': answer}), 200
                except requests.exceptions.RequestException as fallback_req_e:
                    print(f"API Request Error for fallback: {fallback_req_e}")
                    return jsonify({'answer': f'API Error: {str(fallback_req_e)}. Check your API key, model permissions, and request format.'}), 500
            else:
                return jsonify({'answer': f'API Error: {str(req_e)}. Check your API key, model permissions, and request format.'}), 500
    except Exception as e: 
        print(f"General Error for ask: {e}")
        return jsonify({'answer': f'Error: {str(e)}'}), 500

@app.route('/generate_questions', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_questions():
    if request.method == 'OPTIONS': return '', 200
    if not API_KEY: return jsonify({'questions': [], 'message': 'API key not configured.'}), 500
    try:
        data = request.get_json()
        document_text = data.get('documentText', '')
        department = data.get('department', 'general study')

        if not document_text.strip():
            return jsonify({'questions': [], 'message': 'No document text provided for question generation.'}), 400

        # Enhanced prompt for question generation
        prompt = f"""As UniStudy AI, your task is to act as a brilliant question generator, helping students assess their understanding.
        Generate 3-5 concise, insightful, and highly relevant questions that a student might ask based on the following text from the {department} department.
        These questions should encourage critical thinking and comprehension.
        Provide the questions as a JSON array of strings. Example: ["Question 1?", "Question 2?"].
        Ensure the questions are clear, directly related to the provided text, and varied in scope to cover different aspects of the content.
        Text:
        {document_text}
        """
        
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "ARRAY",
                    "items": {"type": "STRING"}
                },
                "maxOutputTokens": 200 # Limit output to keep questions concise
            }
        }

        response = requests.post(f"{GEMINI_TEXT_API_URL}?key={API_KEY}", headers={'Content-Type': 'application/json'}, json=payload)
        response.raise_for_status()
        result = response.json()
        
        generated_questions = []
        if result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts'):
            try:
                # The response is already JSON, so parse it directly
                json_string = result['candidates'][0]['content']['parts'][0]['text']
                generated_questions = json.loads(json_string)
                if not isinstance(generated_questions, list):
                    generated_questions = [] # Ensure it's a list
            except json.JSONDecodeError as e:
                print(f"JSON decode error in question generation: {e}")
                generated_questions = []
        
        return jsonify({'questions': generated_questions, 'message': 'Questions generated successfully.'}), 200

    except requests.exceptions.RequestException as req_e: 
        print(f"API Request Error for generate_questions: {req_e}")
        # Added more specific error message for 403
        if req_e.response and req_e.response.status_code == 403:
            return jsonify({'questions': [], 'message': f'API Error: 403 Forbidden. Your API key might not have permissions for this model or API. Please check Google Cloud Console.'}), 500
        else:
            return jsonify({'questions': [], 'message': f'API Error: {str(req_e)}. Check your API key and permissions.'}), 500
    except Exception as e: 
        print(f"General Error for generate_questions: {e}")
        return jsonify({'questions': [], 'message': f'Error: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='localhost', port=5000)
